import streamlit as st
from groq import Groq
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io, os, smtplib, re
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
from datetime import date
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np

# PDF
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.units import cm
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
                                 TableStyle, Image as RLImage, HRFlowable,
                                 PageBreak, KeepTogether)

# ══════════════════════════════════════════════════════════════
#  CONFIG
# ══════════════════════════════════════════════════════════════
GMAIL_USER      = "Wijdan.psyc@gmail.com"
GMAIL_PASS      = "rias eeul lyuu stce"
RECIPIENT_EMAIL = "Wijdan.psyc@gmail.com"
LOGO_FILE       = "logo.png"

CLINIC_BLUE_RGB = RGBColor(0x8B, 0x73, 0x55)
DARK_RGB        = RGBColor(0x1C, 0x19, 0x17)
WARM_RGB        = RGBColor(0x8B, 0x73, 0x55)
ACCENT_RGB      = RGBColor(0xC4, 0x95, 0x6A)

PDF_WARM   = colors.HexColor('#8B7355')
PDF_DARK   = colors.HexColor('#1C1917')
PDF_CREAM  = colors.HexColor('#F7F3EE')
PDF_BORDER = colors.HexColor('#DDD5C8')
PDF_ACCENT = colors.HexColor('#C4956A')
PDF_RED    = colors.HexColor('#C62828')
PDF_ORANGE = colors.HexColor('#F57C00')
PDF_YELLOW = colors.HexColor('#FBC02D')
PDF_GREEN  = colors.HexColor('#388E3C')
PDF_HEADER = colors.HexColor('#2D2926')

# ══════════════════════════════════════════════════════════════
#  CBCL/6-18  113 ITEMS  (English + Arabic)
# ══════════════════════════════════════════════════════════════
ITEMS_EN = [
    "Acts too young for his/her age",                                          # 1
    "Drinks alcohol without parents' approval",                                # 2
    "Argues a lot",                                                            # 3
    "Fails to finish things he/she starts",                                    # 4
    "There is very little he/she enjoys",                                      # 5
    "Bowel movements outside toilet",                                          # 6
    "Bragging, boasting",                                                      # 7
    "Can't concentrate, can't pay attention for long",                         # 8
    "Can't get his/her mind off certain thoughts; obsessions",                 # 9
    "Can't sit still, restless, or hyperactive",                              # 10
    "Clings to adults or too dependent",                                      # 11
    "Complains of loneliness",                                                # 12
    "Confused or seems to be in a fog",                                       # 13
    "Cries a lot",                                                            # 14
    "Cruel to animals",                                                       # 15
    "Cruelty, bullying, or meanness to others",                               # 16
    "Daydreams or gets lost in his/her thoughts",                             # 17
    "Deliberately harms self or attempts suicide",                            # 18
    "Demands a lot of attention",                                             # 19
    "Destroys his/her own things",                                            # 20
    "Destroys things belonging to his/her family or others",                  # 21
    "Disobedient at home",                                                    # 22
    "Disobedient at school",                                                  # 23
    "Doesn't eat well",                                                       # 24
    "Doesn't get along with other kids",                                      # 25
    "Doesn't seem to feel guilty after misbehaving",                          # 26
    "Easily jealous",                                                         # 27
    "Breaks rules at home, school, or elsewhere",                             # 28
    "Fears certain animals, situations, or places other than school",         # 29
    "Fears going to school",                                                  # 30
    "Fears he/she might think or do something bad",                           # 31
    "Feels he/she has to be perfect",                                         # 32
    "Feels or complains that no one loves him/her",                           # 33
    "Feels others are out to get him/her",                                    # 34
    "Feels worthless or inferior",                                            # 35
    "Gets hurt a lot, accident-prone",                                        # 36
    "Gets in many fights",                                                    # 37
    "Gets teased a lot",                                                      # 38
    "Hangs around with others who get in trouble",                            # 39
    "Hears sounds or voices that aren't there",                               # 40
    "Impulsive or acts without thinking",                                     # 41
    "Would rather be alone than with others",                                 # 42
    "Lying or cheating",                                                      # 43
    "Bites fingernails",                                                      # 44
    "Nervous, highstrung, or tense",                                          # 45
    "Nervous movements or twitching",                                         # 46
    "Nightmares",                                                             # 47
    "Not liked by other kids",                                                # 48
    "Constipated, doesn't move bowels",                                       # 49
    "Too fearful or anxious",                                                 # 50
    "Feels dizzy or lightheaded",                                             # 51
    "Feels too guilty",                                                       # 52
    "Overeating",                                                             # 53
    "Overtired without good reason",                                          # 54
    "Overweight",                                                             # 55
    "Physical problems — aches or pains (not stomach or headaches)",          # 56a→56
    "Physical problems — headaches",                                          # 56b→57
    "Physical problems — nausea, feels sick",                                 # 56c→58
    "Physical problems — stomachaches",                                       # 56f→59
    "Physical problems — vomiting, throwing up",                              # 56g→60
    "Physically attacks people",                                              # 57→61
    "Poor school work",                                                       # 61→62
    "Poorly coordinated or clumsy",                                           # 62→63
    "Prefers being with older kids",                                          # 63→64
    "Prefers being with younger kids",                                        # 64→65
    "Refuses to talk",                                                        # 65→66
    "Repeats certain acts over and over; compulsions",                        # 66→67
    "Runs away from home",                                                    # 67→68
    "Screams a lot",                                                          # 68→69
    "Secretive, keeps things to self",                                        # 69→70
    "Sees things that aren't there",                                          # 70→71
    "Self-conscious or easily embarrassed",                                   # 71→72
    "Sets fires",                                                             # 72→73
    "Too shy or timid",                                                       # 75→74
    "Sleeps less than most kids",                                             # 76→75
    "Sleeps more than most kids",                                             # 77→76
    "Inattentive or easily distracted",                                       # 78→77
    "Stares blankly",                                                         # 80→78
    "Steals at home",                                                         # 81→79
    "Steals outside the home",                                                # 82→80
    "Stores up too many things he/she doesn't need",                          # 83→81
    "Strange behavior",                                                       # 84→82
    "Strange ideas",                                                          # 85→83
    "Stubborn, sullen, or irritable",                                         # 86→84
    "Sudden changes in mood or feelings",                                     # 87→85
    "Sulks a lot",                                                            # 88→86
    "Suspicious",                                                             # 89→87
    "Swearing or obscene language",                                           # 90→88
    "Talks about killing self",                                               # 91→89
    "Talks too much",                                                         # 93→90
    "Teases a lot",                                                           # 94→91
    "Temper tantrums or hot temper",                                          # 95→92
    "Threatens people",                                                       # 97→93
    "Truancy, skips school",                                                  # 101→94
    "Underactive, slow moving, or lacks energy",                              # 102→95
    "Unhappy, sad, or depressed",                                             # 103→96
    "Uses drugs for nonmedical purposes",                                     # 105→97
    "Vandalism",                                                              # 106→98
    "Wets self during the day",                                               # 107→99
    "Wets the bed",                                                           # 108→100
    "Whining",                                                                # 109→101
    "Withdrawn, doesn't get involved with others",                            # 111→102
    "Worries",                                                                # 112→103
]

ITEMS_AR = [
    "يتصرف بشكل أصغر من سنه",                                                # 1
    "يشرب الكحول دون موافقة والديه",                                          # 2
    "يجادل كثيراً",                                                           # 3
    "لا ينهي ما يبدأه",                                                       # 4
    "لا يوجد شيء يستمتع به",                                                  # 5
    "يتبرز خارج المرحاض",                                                     # 6
    "يتفاخر ويتباهى",                                                         # 7
    "لا يستطيع التركيز أو الانتباه لفترة طويلة",                             # 8
    "لا يستطيع إخراج أفكار معينة من ذهنه (وساوس)",                           # 9
    "لا يستطيع الجلوس بهدوء، مضطرب أو مفرط النشاط",                         # 10
    "يتعلق بالبالغين أو يعتمد عليهم كثيراً",                                # 11
    "يشتكي من الوحدة",                                                        # 12
    "مرتبك أو يبدو في ضبابية",                                               # 13
    "يبكي كثيراً",                                                            # 14
    "قاسٍ على الحيوانات",                                                     # 15
    "قسوة أو تنمر أو شراسة مع الآخرين",                                      # 16
    "يحلم أحلام اليقظة أو يضيع في أفكاره",                                  # 17
    "يؤذي نفسه عمداً أو يحاول الانتحار",                                     # 18
    "يطلب الكثير من الاهتمام",                                               # 19
    "يدمر أشياءه الخاصة",                                                    # 20
    "يدمر أشياء تعود لعائلته أو للآخرين",                                    # 21
    "عاصٍ في المنزل",                                                         # 22
    "عاصٍ في المدرسة",                                                        # 23
    "لا يأكل بشكل جيد",                                                      # 24
    "لا يتعامل بشكل جيد مع الأطفال الآخرين",                                # 25
    "لا يشعر بالذنب بعد سوء التصرف",                                         # 26
    "يغار بسهولة",                                                            # 27
    "يكسر القواعد في المنزل أو المدرسة أو غيرها",                           # 28
    "يخاف من حيوانات أو مواقف أو أماكن معينة (غير المدرسة)",                # 29
    "يخاف من الذهاب إلى المدرسة",                                            # 30
    "يخاف أن يفكر أو يفعل شيئاً سيئاً",                                     # 31
    "يشعر أنه يجب أن يكون كاملاً",                                           # 32
    "يشعر أو يشتكي من أن لا أحد يحبه",                                      # 33
    "يشعر أن الآخرين يريدون الإيقاع به",                                     # 34
    "يشعر بعدم القيمة أو الدونية",                                           # 35
    "يُصاب كثيراً، عرضة للحوادث",                                            # 36
    "يشارك في مشاجرات كثيرة",                                                # 37
    "يتعرض للسخرية كثيراً",                                                   # 38
    "يتردد مع من يسبب المشاكل",                                               # 39
    "يسمع أصواتاً أو أشخاصاً غير موجودين",                                   # 40
    "متهور أو يتصرف دون تفكير",                                              # 41
    "يفضل أن يكون وحيداً على أن يكون مع الآخرين",                           # 42
    "يكذب أو يغش",                                                            # 43
    "يعض أظافره",                                                             # 44
    "عصبي أو متوتر",                                                          # 45
    "حركات عصبية أو ارتعاش",                                                 # 46
    "كوابيس",                                                                 # 47
    "لا يُحبه الأطفال الآخرون",                                               # 48
    "إمساك، لا تتحرك الأمعاء",                                               # 49
    "خائف جداً أو قلق",                                                       # 50
    "يشعر بالدوار أو خفة الرأس",                                             # 51
    "يشعر بذنب مفرط",                                                         # 52
    "يأكل بإفراط",                                                            # 53
    "مرهق جداً دون سبب واضح",                                                # 54
    "يعاني من زيادة الوزن",                                                  # 55
    "مشاكل جسدية — آلام وأوجاع (غير المعدة والرأس)",                         # 56
    "مشاكل جسدية — صداع",                                                    # 57
    "مشاكل جسدية — غثيان",                                                   # 58
    "مشاكل جسدية — آلام في المعدة",                                          # 59
    "مشاكل جسدية — تقيؤ",                                                    # 60
    "يهاجم الناس جسدياً",                                                    # 61
    "أداؤه الدراسي ضعيف",                                                    # 62
    "سيء التنسيق أو أخرق",                                                   # 63
    "يفضل الكبار عنه سناً",                                                  # 64
    "يفضل الصغار عنه سناً",                                                  # 65
    "يرفض الكلام",                                                            # 66
    "يكرر أفعالاً معينة مراراً (وسواس قهري)",                               # 67
    "يهرب من المنزل",                                                         # 68
    "يصرخ كثيراً",                                                            # 69
    "كتوم، يحتفظ بالأشياء لنفسه",                                            # 70
    "يرى أشياء غير موجودة",                                                  # 71
    "خجول أو محرج بسهولة",                                                   # 72
    "يضرم النيران",                                                           # 73
    "خجول جداً أو متردد",                                                     # 74
    "ينام أقل من معظم الأطفال",                                               # 75
    "ينام أكثر من معظم الأطفال",                                              # 76
    "غير منتبه أو يتشتت بسهولة",                                             # 77
    "يحدق بنظرة فارغة",                                                      # 78
    "يسرق في المنزل",                                                         # 79
    "يسرق خارج المنزل",                                                       # 80
    "يخزن أشياء لا يحتاج إليها",                                             # 81
    "سلوك غريب",                                                              # 82
    "أفكار غريبة",                                                            # 83
    "عنيد، متجهم، أو سريع الانفعال",                                         # 84
    "تغيرات مفاجئة في المزاج أو المشاعر",                                    # 85
    "يتعكر مزاجه كثيراً",                                                    # 86
    "مريب",                                                                   # 87
    "يلفظ كلاماً نابياً أو بذيئاً",                                          # 88
    "يتحدث عن الانتحار",                                                     # 89
    "يتكلم كثيراً",                                                           # 90
    "يستفز كثيراً",                                                           # 91
    "نوبات غضب أو انفعال شديد",                                              # 92
    "يهدد الناس",                                                             # 93
    "يتغيب عن المدرسة بدون إذن",                                             # 94
    "خامل أو بطيء الحركة أو يفتقر إلى الطاقة",                              # 95
    "يشعر بالتعاسة أو الحزن أو الاكتئاب",                                   # 96
    "يستخدم المخدرات لأغراض غير طبية",                                       # 97
    "تخريب ممتلكات الآخرين",                                                 # 98
    "يبلل نفسه أثناء النهار",                                                # 99
    "يبلل فراشه",                                                             # 100
    "يتذمر",                                                                  # 101
    "منسحب، لا يشارك مع الآخرين",                                            # 102
    "يقلق",                                                                   # 103
]

# ══════════════════════════════════════════════════════════════
#  CBCL SUBSCALES & SCORING
#  Items mapped to app item numbers (1-103 as listed above)
#  Source: Achenbach & Rescorla (2001) CBCL/6-18 scoring keys
# ══════════════════════════════════════════════════════════════
# Mapping from CBCL original item numbers to our sequential 1-103 list
ORIG_TO_SEQ = {
    1:1,2:2,3:3,4:4,5:5,6:6,7:7,8:8,9:9,10:10,11:11,12:12,13:13,14:14,
    15:15,16:16,17:17,18:18,19:19,20:20,21:21,22:22,23:23,24:24,25:25,
    26:26,27:27,28:28,29:29,30:30,31:31,32:32,33:33,34:34,35:35,36:36,
    37:37,38:38,39:39,40:40,41:41,42:42,43:43,44:44,45:45,46:46,47:47,
    48:48,49:49,50:50,51:51,52:52,53:53,54:54,55:55,
    # 56a-g → items 56-60 (we use a,b,c,f,g only)
    "56a":56,"56b":57,"56c":58,"56f":59,"56g":60,
    57:61,61:62,62:63,63:64,64:65,65:66,66:67,67:68,68:69,69:70,70:71,
    71:72,72:73,75:74,76:75,77:76,78:77,80:78,81:79,82:80,83:81,84:82,
    85:83,86:84,87:85,88:86,89:87,90:88,91:89,93:90,94:91,95:92,97:93,
    101:94,102:95,103:96,105:97,106:98,107:99,108:100,109:101,111:102,112:103,
}

def o(n):
    """Get sequential item number from original CBCL item number."""
    return ORIG_TO_SEQ.get(n, 0)

SUBSCALES = {
    # ── 8 Syndrome Scales ──
    "ANX": {
        "name_en": "Anxious/Depressed",
        "name_ar": "القلق / الاكتئاب",
        "items": [o(14),o(29),o(30),o(31),o(32),o(33),o(35),o(45),o(50),o(52),o(71),o(91),o(112)],
        "color": "#6A1B9A", "group": "INT",
    },
    "WIT": {
        "name_en": "Withdrawn/Depressed",
        "name_ar": "الانسحاب / الاكتئاب",
        "items": [o(5),o(42),o(65),o(69),o(75),o(102),o(103),o(111)],
        "color": "#1565C0", "group": "INT",
    },
    "SOM": {
        "name_en": "Somatic Complaints",
        "name_ar": "الشكاوى الجسدية",
        "items": [o(47),o(49),o(51),o(54),"56a","56b","56c","56f","56g"],
        "color": "#4E342E", "group": "INT",
    },
    "SOC": {
        "name_en": "Social Problems",
        "name_ar": "المشكلات الاجتماعية",
        "items": [o(11),o(12),o(25),o(27),o(34),o(36),o(38),o(48),o(62),o(64),o(79)],
        "color": "#0277BD", "group": "MIX",
    },
    "THO": {
        "name_en": "Thought Problems",
        "name_ar": "مشكلات التفكير",
        "items": [o(9),o(18),o(40),o(46),o(58),o(59),o(60),o(66),o(70),o(76),o(83),o(84),o(85),o(92),o(100)],
        "color": "#827717", "group": "MIX",
    },
    "ATT": {
        "name_en": "Attention Problems",
        "name_ar": "مشكلات الانتباه",
        "items": [o(1),o(4),o(8),o(10),o(13),o(17),o(41),o(61),o(78)],
        "color": "#E65100", "group": "MIX",
    },
    "RUL": {
        "name_en": "Rule-Breaking Behavior",
        "name_ar": "السلوك الخارج عن القانون",
        "items": [o(2),o(26),o(28),o(39),o(43),o(63),o(67),o(72),o(73),o(81),o(82),o(90),o(96),o(99),o(101),o(105),o(106)],
        "color": "#B71C1C", "group": "EXT",
    },
    "AGG": {
        "name_en": "Aggressive Behavior",
        "name_ar": "السلوك العدواني",
        "items": [o(3),o(16),o(19),o(20),o(21),o(22),o(23),o(37),o(57),o(68),o(86),o(87),o(88),o(89),o(94),o(95),o(97),o(104)],
        "color": "#C62828", "group": "EXT",
    },
    # ── Broadband ──
    "INT": {
        "name_en": "Internalizing Problems",
        "name_ar": "المشكلات الداخلية",
        "items": [],  # computed as ANX+WIT+SOM
        "color": "#4527A0", "group": "BROAD",
    },
    "EXT": {
        "name_en": "Externalizing Problems",
        "name_ar": "المشكلات الخارجية",
        "items": [],  # computed as RUL+AGG
        "color": "#BF360C", "group": "BROAD",
    },
    "TOT": {
        "name_en": "Total Problems",
        "name_ar": "إجمالي المشكلات",
        "items": [],  # all syndrome items combined
        "color": "#1B5E20", "group": "BROAD",
    },
    # ── DSM-Oriented Scales ──
    "DEP": {
        "name_en": "DSM: Depressive Problems",
        "name_ar": "مشكلات الاكتئاب DSM",
        "items": [o(5),o(14),o(18),o(24),o(35),o(52),o(54),o(76),o(77),o(91),o(100),o(102),o(103)],
        "color": "#283593", "group": "DSM",
    },
    "ANX_D": {
        "name_en": "DSM: Anxiety Problems",
        "name_ar": "مشكلات القلق DSM",
        "items": [o(11),o(29),o(30),o(31),o(45),o(47),o(50),o(71),o(112)],
        "color": "#6A1B9A", "group": "DSM",
    },
    "SOM_D": {
        "name_en": "DSM: Somatic Problems",
        "name_ar": "المشكلات الجسدية DSM",
        "items": [o("56a"),o("56b"),o("56c"),o("56f"),o("56g")],
        "color": "#4E342E", "group": "DSM",
    },
    "ADHD_D": {
        "name_en": "DSM: ADHD",
        "name_ar": "اضطراب ADHD DSM",
        "items": [o(4),o(8),o(10),o(41),o(78),o(93),o(104)],
        "color": "#E65100", "group": "DSM",
    },
    "ODD": {
        "name_en": "DSM: Oppositional Defiant",
        "name_ar": "اضطراب التمرد والعصيان DSM",
        "items": [o(3),o(22),o(23),o(86),o(95)],
        "color": "#AD1457", "group": "DSM",
    },
    "CON": {
        "name_en": "DSM: Conduct Problems",
        "name_ar": "اضطراب السلوك DSM",
        "items": [o(15),o(16),o(21),o(26),o(28),o(37),o(39),o(43),o(57),o(67),o(72),o(81),o(82),o(90),o(97),o(101),o(106)],
        "color": "#BF360C", "group": "DSM",
    },
}

# Fix SOM subscale — replace string keys with actual seq numbers
for k in ["SOM","SOM_D"]:
    fixed = []
    for it in SUBSCALES[k]["items"]:
        if isinstance(it, str):
            fixed.append(ORIG_TO_SEQ.get(it, 0))
        else:
            fixed.append(it)
    SUBSCALES[k]["items"] = [x for x in fixed if x > 0]

# Broadband: build item lists
SUBSCALES["INT"]["items"] = list(set(
    SUBSCALES["ANX"]["items"] + SUBSCALES["WIT"]["items"] + SUBSCALES["SOM"]["items"]
))
SUBSCALES["EXT"]["items"] = list(set(
    SUBSCALES["RUL"]["items"] + SUBSCALES["AGG"]["items"]
))
SUBSCALES["TOT"]["items"] = list(set(
    SUBSCALES["ANX"]["items"] + SUBSCALES["WIT"]["items"] + SUBSCALES["SOM"]["items"] +
    SUBSCALES["SOC"]["items"] + SUBSCALES["THO"]["items"] + SUBSCALES["ATT"]["items"] +
    SUBSCALES["RUL"]["items"] + SUBSCALES["AGG"]["items"]
))

# Remove zeros
for k in SUBSCALES:
    SUBSCALES[k]["items"] = [x for x in SUBSCALES[k]["items"] if x > 0]

# Approximate published norms (Achenbach & Rescorla 2001, combined gender)
NORMS = {
    "ANX":  (3.5,  3.8),
    "WIT":  (2.2,  2.5),
    "SOM":  (1.8,  2.4),
    "SOC":  (2.5,  2.9),
    "THO":  (1.2,  2.1),
    "ATT":  (4.0,  3.5),
    "RUL":  (1.8,  2.6),
    "AGG":  (5.2,  5.5),
    "INT":  (7.5,  6.8),
    "EXT":  (7.0,  7.5),
    "TOT":  (24.0, 18.0),
    "DEP":  (2.8,  3.2),
    "ANX_D":(1.6,  2.2),
    "SOM_D":(0.8,  1.5),
    "ADHD_D":(2.5, 2.8),
    "ODD":  (2.0,  2.5),
    "CON":  (1.2,  2.0),
}

SCALE_ORDER = ["ANX","WIT","SOM","SOC","THO","ATT","RUL","AGG",
               "INT","EXT","TOT","DEP","ANX_D","SOM_D","ADHD_D","ODD","CON"]

def get_level_en(t):
    if t >= 70:   return "Clinical Range"
    elif t >= 65: return "Borderline Clinical"
    elif t >= 60: return "Worth Monitoring"
    elif t >= 40: return "Normal Range"
    else:         return "Below Average"

def get_level_ar(t):
    if t >= 70:   return "النطاق الإكلينيكي"
    elif t >= 65: return "حدي الإكلينيكي"
    elif t >= 60: return "يستحق المتابعة"
    elif t >= 40: return "ضمن المتوسط الطبيعي"
    else:         return "أقل من المتوسط"

def get_bar_color(t):
    if t >= 70:   return "#D32F2F"
    elif t >= 65: return "#F57C00"
    elif t >= 60: return "#FBC02D"
    elif t >= 40: return "#388E3C"
    else:         return "#1976D2"

def raw_to_t(raw, key):
    mean, sd = NORMS[key]
    if sd == 0: return 50
    return max(20, min(90, round(50 + 10 * (raw - mean) / sd)))

def compute_scores(responses):
    results = {}
    for key, info in SUBSCALES.items():
        raw = sum(responses.get(i, 0) for i in info["items"] if i > 0)
        max_raw = len(info["items"]) * 2
        results[key] = {
            "raw": raw,
            "t": raw_to_t(raw, key),
            "max_raw": max_raw,
        }
    return results

CRITICAL_ITEMS = [18, 40, 61, 68, 73, 89, 97, 99, 100, 107]  # original CBCL numbers → seq

# ══════════════════════════════════════════════════════════════
#  CHARTS
# ══════════════════════════════════════════════════════════════
def make_syndrome_chart(scores):
    """Bar chart for 8 syndrome scales."""
    keys   = ["ANX","WIT","SOM","SOC","THO","ATT","RUL","AGG"]
    labels = [SUBSCALES[k]["name_en"] for k in keys]
    t_vals = [scores[k]["t"] for k in keys]
    bar_colors = [get_bar_color(t) for t in t_vals]

    fig, ax = plt.subplots(figsize=(11, 6))
    fig.patch.set_facecolor('#F7F3EE'); ax.set_facecolor('#F7F3EE')
    y_pos = np.arange(len(labels))
    bars = ax.barh(y_pos, t_vals, color=bar_colors, edgecolor='white', linewidth=0.8, height=0.6)

    for xv, lbl, col in [(60,'T=60','#FBC02D'),(65,'T=65','#F57C00'),(70,'T=70','#D32F2F')]:
        ax.axvline(x=xv, color=col, linestyle='--', linewidth=1.2, alpha=0.8, label=lbl)

    ax.axvspan(70, 95, alpha=0.07, color='#D32F2F')
    ax.axvspan(65, 70, alpha=0.06, color='#F57C00')

    for bar_, val in zip(bars, t_vals):
        ax.text(bar_.get_width()+0.6, bar_.get_y()+bar_.get_height()/2,
                f"{val}", va='center', ha='left', fontsize=9, fontweight='bold', color='#1C1917')

    # Labels: Internalizing / Externalizing brackets
    n_int, n_ext = 3, 2
    ax.axhline(y=n_int-0.5, color='#1565C0', linestyle=':', lw=1, alpha=0.6)
    ax.axhline(y=n_int+3-0.5, color='#C62828', linestyle=':', lw=1, alpha=0.6)
    ax.text(92, 1, "INT", fontsize=8, color='#1565C0', va='center', style='italic')
    ax.text(92, 6.5, "EXT", fontsize=8, color='#C62828', va='center', style='italic')

    ax.set_yticks(y_pos); ax.set_yticklabels(labels, fontsize=10, fontfamily='DejaVu Sans')
    ax.set_xlim(20, 95)
    ax.set_xlabel('T-Score', fontsize=11, fontweight='bold', color='#1C1917')
    ax.set_title("CBCL/6-18 — Syndrome Scale T-Score Profile",
                 fontsize=13, fontweight='bold', color='#1C1917', pad=12)
    ax.legend(loc='lower right', fontsize=8.5, framealpha=0.7)
    ax.spines['top'].set_visible(False); ax.spines['right'].set_visible(False)
    ax.grid(axis='x', linestyle=':', alpha=0.4)
    plt.tight_layout()
    buf = io.BytesIO(); plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    plt.close(fig); buf.seek(0)
    return buf.read()

def make_broadband_chart(scores):
    """Bar chart: Internalizing, Externalizing, Total Problems."""
    keys   = ["INT","EXT","TOT"]
    labels = ["Internalizing", "Externalizing", "Total Problems"]
    t_vals = [scores[k]["t"] for k in keys]
    raw_vals = [scores[k]["raw"] for k in keys]
    bar_colors = [get_bar_color(t) for t in t_vals]

    fig, ax = plt.subplots(figsize=(7, 4))
    fig.patch.set_facecolor('#F7F3EE'); ax.set_facecolor('#F7F3EE')
    x_pos = np.arange(len(labels))
    bars = ax.bar(x_pos, t_vals, color=bar_colors, edgecolor='white', linewidth=1.2, width=0.55)

    for xv, lbl, col in [(65,'T=65','#F57C00'),(70,'T=70','#D32F2F')]:
        ax.axhline(y=xv, color=col, linestyle='--', linewidth=1.2, alpha=0.8, label=lbl)

    for bar_, t, raw in zip(bars, t_vals, raw_vals):
        ax.text(bar_.get_x()+bar_.get_width()/2, bar_.get_height()+0.8,
                f"T={t}\nRaw={raw}", ha='center', va='bottom',
                fontsize=9, fontweight='bold', color='#1C1917')

    ax.set_xticks(x_pos); ax.set_xticklabels(labels, fontsize=11, fontfamily='DejaVu Sans')
    ax.set_ylim(20, 95)
    ax.set_ylabel('T-Score', fontsize=10)
    ax.set_title("Broadband Problem Scales", fontsize=12, fontweight='bold', color='#1C1917', pad=10)
    ax.legend(fontsize=9, framealpha=0.7)
    ax.spines['top'].set_visible(False); ax.spines['right'].set_visible(False)
    ax.grid(axis='y', linestyle=':', alpha=0.4)
    plt.tight_layout()
    buf = io.BytesIO(); plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    plt.close(fig); buf.seek(0)
    return buf.read()

def make_dsm_chart(scores):
    """Bar chart for 6 DSM-oriented scales."""
    keys   = ["DEP","ANX_D","SOM_D","ADHD_D","ODD","CON"]
    labels = ["Depressive", "Anxiety", "Somatic", "ADHD", "ODD", "Conduct"]
    t_vals = [scores[k]["t"] for k in keys]
    bar_colors = [get_bar_color(t) for t in t_vals]

    fig, ax = plt.subplots(figsize=(10, 5))
    fig.patch.set_facecolor('#F7F3EE'); ax.set_facecolor('#F7F3EE')
    x_pos = np.arange(len(labels))
    bars = ax.bar(x_pos, t_vals, color=bar_colors, edgecolor='white', linewidth=1.0, width=0.6)

    for xv, lbl, col in [(65,'T=65 (Borderline)','#F57C00'),(70,'T=70 (Clinical)','#D32F2F')]:
        ax.axhline(y=xv, color=col, linestyle='--', linewidth=1.2, alpha=0.8, label=lbl)

    for bar_, val in zip(bars, t_vals):
        ax.text(bar_.get_x()+bar_.get_width()/2, bar_.get_height()+0.6,
                str(val), ha='center', va='bottom', fontsize=10, fontweight='bold', color='#1C1917')

    ax.set_xticks(x_pos); ax.set_xticklabels(labels, fontsize=11, fontfamily='DejaVu Sans')
    ax.set_ylim(20, 95)
    ax.set_ylabel('T-Score', fontsize=10)
    ax.set_title("CBCL/6-18 — DSM-Oriented Scale Scores",
                 fontsize=12, fontweight='bold', color='#1C1917', pad=10)
    ax.legend(fontsize=9, framealpha=0.7)
    ax.spines['top'].set_visible(False); ax.spines['right'].set_visible(False)
    ax.grid(axis='y', linestyle=':', alpha=0.4)
    plt.tight_layout()
    buf = io.BytesIO(); plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    plt.close(fig); buf.seek(0)
    return buf.read()

def make_pie_chart(responses):
    counts = [0, 0, 0]
    for v in responses.values(): 
        if v in [0,1,2]: counts[v] += 1
    labels  = ['0 — Not True','1 — Somewhat True','2 — Very True']
    clrs    = ['#388E3C','#FBC02D','#D32F2F']
    fig, ax = plt.subplots(figsize=(6, 4.5))
    fig.patch.set_facecolor('#F7F3EE')
    wedges, texts, autotexts = ax.pie(
        counts, labels=labels, colors=clrs,
        autopct='%1.0f%%', startangle=90,
        wedgeprops={'edgecolor':'white','linewidth':1.5}
    )
    for at in autotexts: at.set_fontsize(9); at.set_fontweight('bold')
    ax.set_title('Response Distribution', fontsize=11, fontweight='bold', color='#1C1917')
    plt.tight_layout()
    buf = io.BytesIO(); plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    plt.close(fig); buf.seek(0)
    return buf.read()

# ══════════════════════════════════════════════════════════════
#  GROQ REPORTS
# ══════════════════════════════════════════════════════════════
def _score_block_en(scores):
    lines = []
    groups = [
        ("SYNDROME SCALES", ["ANX","WIT","SOM","SOC","THO","ATT","RUL","AGG"]),
        ("BROADBAND SCALES", ["INT","EXT","TOT"]),
        ("DSM-ORIENTED SCALES", ["DEP","ANX_D","SOM_D","ADHD_D","ODD","CON"]),
    ]
    for grp, keys in groups:
        lines.append(f"\n  {grp}:")
        for k in keys:
            s = scores[k]
            lines.append(f"    {SUBSCALES[k]['name_en']}: Raw={s['raw']}, T={s['t']} — {get_level_en(s['t'])}")
    return "\n".join(lines)

def _score_block_ar(scores):
    lines = []
    groups = [
        ("مقاييس المتلازمات", ["ANX","WIT","SOM","SOC","THO","ATT","RUL","AGG"]),
        ("المقاييس الشاملة", ["INT","EXT","TOT"]),
        ("مقاييس موجهة DSM", ["DEP","ANX_D","SOM_D","ADHD_D","ODD","CON"]),
    ]
    for grp, keys in groups:
        lines.append(f"\n  {grp}:")
        for k in keys:
            s = scores[k]
            lines.append(f"    {SUBSCALES[k]['name_ar']}: خام={s['raw']}, تائي={s['t']} — {get_level_ar(s['t'])}")
    return "\n".join(lines)

def generate_report_en(child_name, age, gender, rater, scores):
    elevated_syn  = [k for k in ["ANX","WIT","SOM","SOC","THO","ATT","RUL","AGG"] if scores[k]["t"] >= 65]
    elevated_dsm  = [k for k in ["DEP","ANX_D","SOM_D","ADHD_D","ODD","CON"]     if scores[k]["t"] >= 65]
    int_t   = scores["INT"]["t"]; ext_t = scores["EXT"]["t"]; tot_t = scores["TOT"]["t"]
    prompt = f"""You are a licensed child clinical psychologist writing a professional CBCL/6-18 assessment report.

CHILD: {child_name} | AGE: {age} | GENDER: {gender} | RATER: {rater}
ASSESSMENT: Child Behavior Checklist for Ages 6-18 (CBCL/6-18) — Achenbach & Rescorla (2001)
DATE: {date.today().strftime('%B %d, %Y')}
RATING PERIOD: Past 6 months

T-SCORES (T≥70 = Clinical Range; T=65-69 = Borderline Clinical; T<65 = Normal):
{_score_block_en(scores)}

ELEVATED SYNDROME SCALES (T≥65): {', '.join(SUBSCALES[k]['name_en'] for k in elevated_syn) if elevated_syn else 'None'}
ELEVATED DSM SCALES (T≥65): {', '.join(SUBSCALES[k]['name_en'] for k in elevated_dsm) if elevated_dsm else 'None'}
Internalizing T={int_t} | Externalizing T={ext_t} | Total Problems T={tot_t}

RULES:
- Do NOT diagnose. Findings as hypotheses only.
- Formal clinical language. No markdown (**, ##, ---).
- Section titles: ALL CAPS numbered.
- Clearly distinguish Internalizing vs Externalizing problem patterns.
- Reference DSM-5 diagnostic categories as possibilities, not conclusions.

REPORT STRUCTURE:

CHILD BEHAVIOR CHECKLIST (CBCL/6-18) — CLINICAL REPORT
Child | {child_name}
Age | {age}  |  Gender | {gender}
Rater | {rater}  |  Date | {date.today().strftime('%B %d, %Y')}
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

CLINICAL SUMMARY
4–6 sentences: broadband profile (Internalizing T={int_t}, Externalizing T={ext_t}, Total T={tot_t}),
most elevated syndrome and DSM scales, overall clinical significance.

1. REFERRAL AND ASSESSMENT OVERVIEW
CBCL/6-18 description, purpose, Achenbach ASEBA system, 113-item format (0-1-2 scale),
rating period (past 6 months), who completed it.

2. INTERNALIZING PROBLEMS PROFILE
Interpret Internalizing T={int_t}. Detailed analysis of Anxious/Depressed, Withdrawn/Depressed,
and Somatic Complaints syndrome scales. Behavioral correlates and clinical significance.

3. EXTERNALIZING PROBLEMS PROFILE
Interpret Externalizing T={ext_t}. Detailed analysis of Rule-Breaking Behavior and Aggressive
Behavior syndrome scales. Behavioral correlates. School and family implications.

4. MIXED PROBLEM SCALES
Interpret Social Problems, Thought Problems, Attention Problems — these fall between
Internalizing and Externalizing domains. Clinical significance for each.

5. DSM-ORIENTED SCALE ANALYSIS
For each DSM scale (Depressive, Anxiety, Somatic, ADHD, ODD, Conduct Problems):
T-score, clinical range classification, diagnostic considerations as hypotheses.

6. TOTAL PROBLEMS AND OVERALL SEVERITY
Interpret Total Problems T={tot_t}. Overall burden of behavioral and emotional problems.
Compare Internalizing vs Externalizing dominance.

7. CRITICAL ITEMS AND NOTABLE RESPONSES
Note any clinically significant individual items that warrant immediate attention
(self-harm, suicidal ideation, fire setting, running away, substance use).

8. STRENGTHS AND PROTECTIVE FACTORS
Scales in normal range as relative strengths. Areas of adaptive functioning.

9. INTEGRATED CLINICAL IMPRESSIONS
Synthesize the full profile. Primary diagnostic hypotheses (NOT formal diagnoses).
Differential considerations based on scale pattern.

10. RECOMMENDATIONS
10-12 specific evidence-based recommendations:
a) Mental health intervention priorities
b) School accommodations and support
c) Parenting strategies / family therapy
d) Further evaluation needs
e) Medical consultation if indicated
f) Community support and resources

11. SUMMARY
One paragraph for clinical records:
"The CBCL/6-18 was completed by {rater} for {child_name} (age {age}, gender {gender}).
Results indicate..."
"""
    client = Groq(api_key=st.secrets["GROQ_API_KEY"])
    r = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[{"role":"user","content":prompt}],
        max_tokens=4000
    )
    return r.choices[0].message.content.strip()

def generate_report_ar(child_name, age, gender, rater, scores):
    elevated_syn = [SUBSCALES[k]["name_ar"] for k in ["ANX","WIT","SOM","SOC","THO","ATT","RUL","AGG"] if scores[k]["t"] >= 65]
    elevated_dsm = [SUBSCALES[k]["name_ar"] for k in ["DEP","ANX_D","SOM_D","ADHD_D","ODD","CON"]     if scores[k]["t"] >= 65]
    int_t  = scores["INT"]["t"]; ext_t = scores["EXT"]["t"]; tot_t = scores["TOT"]["t"]
    prompt = f"""أنت طبيب نفسي للأطفال تكتب تقريراً سريرياً احترافياً لقائمة أكنباك للسلوك الطفلي (CBCL/6-18).

الطفل: {child_name} | السن: {age} | النوع: {gender} | المُقيِّم: {rater}
المقياس: قائمة السلوك الطفلي للأعمار 6-18 (CBCL/6-18) — أكنباك ورسكورلا (2001)
التاريخ: {date.today().strftime('%Y/%m/%d')}
فترة التقييم: الستة أشهر الماضية

الدرجات التائية (T≥70 = النطاق الإكلينيكي؛ T=65-69 = حدي؛ T<65 = طبيعي):
{_score_block_ar(scores)}

المقاييس المرتفعة — متلازمات (T≥65): {', '.join(elevated_syn) if elevated_syn else 'لا يوجد'}
المقاييس المرتفعة — DSM (T≥65): {', '.join(elevated_dsm) if elevated_dsm else 'لا يوجد'}
المشكلات الداخلية T={int_t} | المشكلات الخارجية T={ext_t} | إجمالي المشكلات T={tot_t}

قواعد صارمة:
- لا تضع تشخيصاً. أشر إلى النتائج كفرضيات تحتاج إلى حكم سريري.
- عربية فصحى رسمية. لا إنجليزية إلا للاختصارات (CBCL, DSM, ADHD, ODD, T-score).
- لا رموز markdown. عناوين الأقسام: أرقام عربية + عناوين واضحة.

هيكل التقرير:

قائمة أكنباك للسلوك الطفلي CBCL/6-18 — التقرير السريري
الطفل | {child_name}
السن | {age}  |  النوع | {gender}
المُقيِّم | {rater}  |  التاريخ | {date.today().strftime('%Y/%m/%d')}
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

ملخص سريري
٤–٦ جمل: الملف الشامل (داخلي T={int_t}, خارجي T={ext_t}, إجمالي T={tot_t})، المقاييس الأكثر ارتفاعاً، الدلالة السريرية الكلية.

١. نظرة عامة على التقييم
وصف CBCL/6-18، نظام ASEBA، الصيغة (113 بنداً، مقياس 0-1-2)، فترة التقييم (6 أشهر)، من أكمل الاستبيان.

٢. ملف المشكلات الداخلية
تفسير المشكلات الداخلية T={int_t}. تحليل مفصّل لمقاييس القلق/الاكتئاب، الانسحاب/الاكتئاب، الشكاوى الجسدية. المظاهر السلوكية والدلالة السريرية.

٣. ملف المشكلات الخارجية
تفسير المشكلات الخارجية T={ext_t}. تحليل مفصّل لمقاييس السلوك الخارج عن القانون والسلوك العدواني. المظاهر السلوكية والتداعيات المدرسية والأسرية.

٤. مقاييس المشكلات المختلطة
تفسير المشكلات الاجتماعية، مشكلات التفكير، مشكلات الانتباه. دلالتها السريرية.

٥. تحليل مقاييس DSM الموجهة
لكل مقياس DSM: درجة تائية، تصنيف النطاق، اعتبارات تشخيصية كفرضيات.

٦. إجمالي المشكلات والشدة العامة
تفسير إجمالي المشكلات T={tot_t}. مقارنة الأنماط الداخلية والخارجية.

٧. البنود الحرجة والاستجابات اللافتة
أي بنود ذات دلالة سريرية فورية (إيذاء النفس، أفكار انتحارية، إضرام النار، الهروب، المخدرات).

٨. نقاط القوة والعوامل الوقائية
المقاييس ضمن المتوسط الطبيعي كنقاط قوة نسبية.

٩. الانطباعات السريرية المتكاملة
تركيب الصورة الكلية. الفرضيات التشخيصية (لا تشخيصات رسمية). الاعتبارات التفريقية.

١٠. التوصيات
١٠-١٢ توصية محددة ومبنية على الأدلة:
أ) أولويات التدخل النفسي
ب) التسهيلات المدرسية والدعم
ج) استراتيجيات الوالدين / الإرشاد الأسري
د) احتياجات التقييم الإضافي
هـ) الإحالة الطبية إن كانت مشارة
و) الموارد المجتمعية والمساندة

١١. الملخص
فقرة واحدة مناسبة للسجلات السريرية.
"أكمل {rater} قائمة CBCL/6-18 للطفل {child_name} (السن {age}، النوع {gender}).
تشير النتائج إلى..."
"""
    client = Groq(api_key=st.secrets["GROQ_API_KEY"])
    r = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[{"role":"user","content":prompt}],
        max_tokens=4000
    )
    return r.choices[0].message.content.strip()

# ══════════════════════════════════════════════════════════════
#  PDF BUILDER  (English)
# ══════════════════════════════════════════════════════════════
def _t_band_color(t):
    if t >= 70: return PDF_RED
    if t >= 65: return PDF_ORANGE
    if t >= 60: return PDF_YELLOW
    return PDF_GREEN

def _t_band_label(t):
    if t >= 70: return "Clinical Range"
    if t >= 65: return "Borderline Clinical"
    if t >= 60: return "Worth Monitoring"
    if t >= 40: return "Normal Range"
    return "Below Average"

def _make_pdf_styles():
    styles = {}
    styles['title']    = ParagraphStyle('title',fontName='Helvetica-Bold',fontSize=16,
                            textColor=PDF_DARK,spaceAfter=4,alignment=TA_CENTER)
    styles['subtitle'] = ParagraphStyle('subtitle',fontName='Helvetica',fontSize=9,
                            textColor=PDF_WARM,spaceAfter=2,alignment=TA_CENTER)
    styles['section']  = ParagraphStyle('section',fontName='Helvetica-Bold',fontSize=11,
                            textColor=PDF_WARM,spaceBefore=14,spaceAfter=4)
    styles['body']     = ParagraphStyle('body',fontName='Helvetica',fontSize=9.5,
                            textColor=PDF_DARK,leading=14,spaceAfter=5)
    styles['small']    = ParagraphStyle('small',fontName='Helvetica',fontSize=8,
                            textColor=PDF_WARM,leading=11)
    styles['bold_body']= ParagraphStyle('bold_body',fontName='Helvetica-Bold',fontSize=9.5,
                            textColor=PDF_DARK,leading=14,spaceAfter=3)
    styles['grp_hdr']  = ParagraphStyle('grp_hdr',fontName='Helvetica-Bold',fontSize=9,
                            textColor=colors.white,leading=12,spaceAfter=0)
    return styles

def build_pdf_report_en(report_text, scores, charts, child_name, age, gender, rater, responses_dict):
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
          leftMargin=2*cm, rightMargin=2*cm, topMargin=2*cm, bottomMargin=2*cm)
    S = _make_pdf_styles()
    W = A4[0] - 4*cm
    story = []

    # Logo
    if os.path.exists(LOGO_FILE):
        try:
            logo = RLImage(LOGO_FILE, width=5*cm, height=2.2*cm); logo.hAlign='CENTER'
            story.append(logo); story.append(Spacer(1,4))
        except: pass

    # Title
    story.append(Paragraph("Child Behavior Checklist — Clinical Report", S['title']))
    story.append(Paragraph("CBCL/6-18 · Achenbach & Rescorla (2001) · ASEBA", S['subtitle']))
    story.append(HRFlowable(width=W, thickness=1, color=PDF_BORDER, spaceAfter=10))

    # Demographics
    demo_data = [
        [Paragraph('<b>Child</b>',S['small']),  Paragraph(child_name or '—',S['body']),
         Paragraph('<b>Age</b>',S['small']),    Paragraph(str(age) or '—',S['body'])],
        [Paragraph('<b>Gender</b>',S['small']), Paragraph(gender or '—',S['body']),
         Paragraph('<b>Rater</b>',S['small']),  Paragraph(rater or '—',S['body'])],
        [Paragraph('<b>Date</b>',S['small']),   Paragraph(date.today().strftime('%B %d, %Y'),S['body']),
         Paragraph('<b>Assessment</b>',S['small']), Paragraph('CBCL/6-18 (113 items, 0–2)',S['body'])],
        [Paragraph('<b>Rating Period</b>',S['small']), Paragraph('Past 6 months',S['body']),
         Paragraph('<b>Scale</b>',S['small']),  Paragraph('0=Not True  1=Somewhat  2=Very True',S['body'])],
    ]
    demo_tbl = Table(demo_data, colWidths=[2.5*cm, 6.2*cm, 2.5*cm, 6.2*cm])
    demo_tbl.setStyle(TableStyle([
        ('BACKGROUND',(0,0),(-1,-1),PDF_CREAM),
        ('BOX',(0,0),(-1,-1),0.5,PDF_BORDER),
        ('INNERGRID',(0,0),(-1,-1),0.3,PDF_BORDER),
        ('VALIGN',(0,0),(-1,-1),'MIDDLE'),
        ('TOPPADDING',(0,0),(-1,-1),5),
        ('BOTTOMPADDING',(0,0),(-1,-1),5),
        ('LEFTPADDING',(0,0),(-1,-1),6),
    ]))
    story.append(KeepTogether([demo_tbl])); story.append(Spacer(1,10))

    # Score summary table
    def scale_table_section(title, keys):
        rows = [[
            Paragraph('<b>Scale</b>',S['small']),
            Paragraph('<b>Raw</b>',S['small']),
            Paragraph('<b>T-Score</b>',S['small']),
            Paragraph('<b>Classification</b>',S['small']),
        ]]
        ts_cmds = [
            ('BACKGROUND',(0,0),(-1,0),PDF_HEADER),
            ('TEXTCOLOR',(0,0),(-1,0),colors.white),
            ('FONTNAME',(0,0),(-1,0),'Helvetica-Bold'),
            ('FONTSIZE',(0,0),(-1,-1),8.5),
            ('BOX',(0,0),(-1,-1),0.5,PDF_BORDER),
            ('INNERGRID',(0,0),(-1,-1),0.3,PDF_BORDER),
            ('VALIGN',(0,0),(-1,-1),'MIDDLE'),
            ('TOPPADDING',(0,0),(-1,-1),4),
            ('BOTTOMPADDING',(0,0),(-1,-1),4),
            ('LEFTPADDING',(0,0),(-1,-1),5),
            ('ROWBACKGROUNDS',(0,1),(-1,-1),[colors.white,PDF_CREAM]),
        ]
        for ri, k in enumerate(keys, start=1):
            s = scores[k]; t = s['t']
            ts_cmds.append(('BACKGROUND',(2,ri),(2,ri),_t_band_color(t)))
            ts_cmds.append(('TEXTCOLOR',(2,ri),(2,ri),colors.black))
            ts_cmds.append(('FONTNAME',(2,ri),(2,ri),'Helvetica-Bold'))
            rows.append([
                Paragraph(SUBSCALES[k]['name_en'],S['body']),
                Paragraph(str(s['raw']),S['body']),
                Paragraph(str(t),S['body']),
                Paragraph(_t_band_label(t),S['body']),
            ])
        tbl = Table(rows, colWidths=[6.5*cm,1.8*cm,1.8*cm,7.3*cm])
        tbl.setStyle(TableStyle(ts_cmds))
        story.append(KeepTogether([
            Paragraph(title, S['section']),
            tbl, Spacer(1,6)
        ]))

    scale_table_section("SYNDROME SCALE SCORES",
                        ["ANX","WIT","SOM","SOC","THO","ATT","RUL","AGG"])
    scale_table_section("BROADBAND SCALE SCORES", ["INT","EXT","TOT"])
    scale_table_section("DSM-ORIENTED SCALE SCORES",
                        ["DEP","ANX_D","SOM_D","ADHD_D","ODD","CON"])

    # Charts
    if charts.get("syndrome"):
        story.append(Paragraph("SYNDROME SCALE T-SCORE PROFILE", S['section']))
        img = RLImage(io.BytesIO(charts["syndrome"]), width=W, height=9*cm)
        story.append(img); story.append(Spacer(1,6))

    if charts.get("broadband") and charts.get("dsm"):
        bb_img = RLImage(io.BytesIO(charts["broadband"]), width=W*0.42, height=7*cm)
        dm_img = RLImage(io.BytesIO(charts["dsm"]),       width=W*0.56, height=7*cm)
        two_col = Table([[bb_img, dm_img]], colWidths=[W*0.44, W*0.56])
        two_col.setStyle(TableStyle([('VALIGN',(0,0),(-1,-1),'MIDDLE'),('LEFTPADDING',(1,0),(1,0),8)]))
        story.append(Paragraph("BROADBAND AND DSM-ORIENTED SCALE PROFILES", S['section']))
        story.append(two_col); story.append(Spacer(1,6))

    if charts.get("pie"):
        story.append(Paragraph("RESPONSE DISTRIBUTION", S['section']))
        pie_img = RLImage(io.BytesIO(charts["pie"]), width=8*cm, height=6*cm)
        pie_img.hAlign = 'LEFT'
        story.append(pie_img); story.append(Spacer(1,6))

    # Narrative
    story.append(HRFlowable(width=W, thickness=0.5, color=PDF_BORDER))
    story.append(Paragraph("CLINICAL NARRATIVE REPORT", S['section']))
    sec_pat = re.compile(r'^\d+\.\s+[A-Z][A-Z\s&/()\-:]+$')
    hdr_words = {"CHILD BEHAVIOR CHECKLIST (CBCL/6-18) — CLINICAL REPORT","CLINICAL SUMMARY"}

    for line in report_text.split('\n'):
        ls = line.strip()
        if not ls: story.append(Spacer(1,4)); continue
        if ls.startswith('━') or ls.startswith('═'):
            story.append(HRFlowable(width=W,thickness=0.4,color=PDF_BORDER,spaceAfter=4)); continue
        is_sec = sec_pat.match(ls) or ls in hdr_words or ls.upper() in hdr_words or ls == "CLINICAL SUMMARY"
        if is_sec: story.append(Paragraph(ls, S['section'])); continue
        if '|' in ls:
            parts = [p.strip() for p in ls.split('|') if p.strip()]
            if len(parts) >= 2:
                skip = [("field","value"),("subscale","raw"),("child",""),("scale","")]
                if (parts[0].lower(), parts[1].lower()) not in skip:
                    row_data = [[Paragraph(p, S['body']) for p in parts]]
                    mini = Table(row_data, colWidths=[W/len(parts)]*len(parts))
                    mini.setStyle(TableStyle([
                        ('BOX',(0,0),(-1,-1),0.3,PDF_BORDER),
                        ('INNERGRID',(0,0),(-1,-1),0.3,PDF_BORDER),
                        ('BACKGROUND',(0,0),(-1,-1),PDF_CREAM),
                        ('TOPPADDING',(0,0),(-1,-1),3),
                        ('BOTTOMPADDING',(0,0),(-1,-1),3),
                        ('LEFTPADDING',(0,0),(-1,-1),5),
                    ]))
                    story.append(KeepTogether([mini])); continue
        story.append(Paragraph(ls, S['body']))

    # Item response table
    story.append(PageBreak())
    story.append(Paragraph("ITEM RESPONSES — FULL RATING TABLE", S['section']))
    story.append(Paragraph(
        "Rating key: 0 = Not True  ·  1 = Somewhat or Sometimes True  ·  2 = Very True or Often True",
        S['small']))
    story.append(Spacer(1,6))

    rating_labels = {0:"0 — Not True", 1:"1 — Somewhat True", 2:"2 — Very True"}
    rating_colors = {0:PDF_GREEN, 1:PDF_YELLOW, 2:PDF_RED}
    item_rows = [[Paragraph('<b>#</b>',S['small']),
                  Paragraph('<b>Item</b>',S['small']),
                  Paragraph('<b>Rating</b>',S['small'])]]
    item_ts = [
        ('BACKGROUND',(0,0),(-1,0),PDF_HEADER),('TEXTCOLOR',(0,0),(-1,0),colors.white),
        ('FONTNAME',(0,0),(-1,0),'Helvetica-Bold'),('FONTSIZE',(0,0),(-1,-1),8),
        ('BOX',(0,0),(-1,-1),0.5,PDF_BORDER),('INNERGRID',(0,0),(-1,-1),0.3,PDF_BORDER),
        ('VALIGN',(0,0),(-1,-1),'MIDDLE'),
        ('TOPPADDING',(0,0),(-1,-1),3),('BOTTOMPADDING',(0,0),(-1,-1),3),
        ('LEFTPADDING',(0,0),(-1,-1),4),
    ]
    for i, item_text in enumerate(ITEMS_EN):
        item_num = i + 1
        val = responses_dict.get(item_num, 0)
        bg = colors.white if i % 2 == 0 else PDF_CREAM
        item_ts.append(('BACKGROUND',(0,item_num),(1,item_num),bg))
        item_ts.append(('BACKGROUND',(2,item_num),(2,item_num),rating_colors[val]))
        item_ts.append(('FONTNAME',(2,item_num),(2,item_num),'Helvetica-Bold'))
        item_rows.append([
            Paragraph(str(item_num),S['small']),
            Paragraph(item_text,S['small']),
            Paragraph(rating_labels[val],S['small']),
        ])
    item_tbl = Table(item_rows, colWidths=[1*cm, 12.5*cm, 4*cm], repeatRows=1)
    item_tbl.setStyle(TableStyle(item_ts))
    story.append(item_tbl)
    story.append(Spacer(1,12))
    story.append(HRFlowable(width=W,thickness=0.5,color=PDF_BORDER))
    story.append(Spacer(1,4))
    story.append(Paragraph(
        "This report is strictly confidential. Scores are based on parent/caregiver rating and should be "
        "interpreted in conjunction with clinical judgment and other assessment data. "
        "CBCL/6-18 T-scores ≥65 are considered clinically significant (Borderline Clinical ≥65; Clinical ≥70). "
        "© 2001 T.M. Achenbach, ASEBA, University of Vermont.",
        S['small']))

    doc.build(story); buf.seek(0)
    return buf

# ══════════════════════════════════════════════════════════════
#  WORD DOC BUILDER  (Arabic)
# ══════════════════════════════════════════════════════════════
def build_word_report_ar(report_text, scores, charts, child_name, age, gender, rater, responses_dict):
    doc = Document()
    for sec_ in doc.sections:
        sec_.top_margin=Cm(2.0); sec_.bottom_margin=Cm(2.0)
        sec_.left_margin=Cm(2.2); sec_.right_margin=Cm(2.2)
    for sec_ in doc.sections:
        sp=sec_._sectPr; pb=OxmlElement('w:pgBorders'); pb.set(qn('w:offsetFrom'),'page')
        for side in ('top','left','bottom','right'):
            b=OxmlElement(f'w:{side}'); b.set(qn('w:val'),'single')
            b.set(qn('w:sz'),'10'); b.set(qn('w:space'),'24')
            b.set(qn('w:color'),'8B7355'); pb.append(b)
        sp.append(pb)
    for sec_ in doc.sections:
        ft=sec_.footer; fp=ft.paragraphs[0] if ft.paragraphs else ft.add_paragraph()
        fp.clear(); fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r_=fp.add_run(); r_.font.size=Pt(9); r_.font.color.rgb=WARM_RGB
        for tag,text in [('begin',None),(None,' PAGE '),('end',None)]:
            if tag:
                el=OxmlElement('w:fldChar'); el.set(qn('w:fldCharType'),tag); r_._r.append(el)
            else:
                it=OxmlElement('w:instrText'); it.text=text; r_._r.append(it)

    def rtl(p):
        pPr=p._p.get_or_add_pPr(); pPr.append(OxmlElement("w:bidi"))
        jc=OxmlElement("w:jc"); jc.set(qn("w:val"),"right"); pPr.append(jc)

    def add_para(text,bold=False,size=11,color=None,space_before=0,space_after=4,keep_next=False,italic=False):
        p=doc.add_paragraph()
        p.paragraph_format.space_before=Pt(space_before); p.paragraph_format.space_after=Pt(space_after)
        if keep_next: p.paragraph_format.keep_with_next=True
        rtl(p)
        r_=p.add_run(text); r_.font.size=Pt(size); r_.font.name="Times New Roman"
        r_.font.bold=bold; r_.font.italic=italic
        if color: r_.font.color.rgb=color
        return p

    def add_section_title(text):
        p=doc.add_paragraph()
        p.paragraph_format.space_before=Pt(12); p.paragraph_format.space_after=Pt(3)
        p.paragraph_format.keep_with_next=True; rtl(p)
        r_=p.add_run(text.strip()); r_.font.size=Pt(12); r_.font.name="Times New Roman"
        r_.font.bold=True; r_.font.color.rgb=WARM_RGB
        pPr=p._p.get_or_add_pPr(); pBdr=OxmlElement('w:pBdr')
        b=OxmlElement('w:bottom'); b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'4')
        b.set(qn('w:space'),'1'); b.set(qn('w:color'),'DDD5C8'); pBdr.append(b); pPr.append(pBdr)

    def make_table_2col():
        t=doc.add_table(rows=0,cols=2); t.style='Table Grid'
        try:
            tPr=t._tbl.tblPr; bv=OxmlElement('w:bidiVisual'); tPr.append(bv)
            tW=OxmlElement('w:tblW'); tW.set(qn('w:w'),'9026'); tW.set(qn('w:type'),'dxa'); tPr.append(tW)
            tg=OxmlElement('w:tblGrid')
            for w in [2500,6526]:
                gc=OxmlElement('w:gridCol'); gc.set(qn('w:w'),str(w)); tg.append(gc)
            t._tbl.insert(0,tg)
        except: pass
        return t

    def add_row(tbl, field, value, header=False):
        row=tbl.add_row()
        trPr=row._tr.get_or_add_trPr()
        cs=OxmlElement('w:cantSplit'); cs.set(qn('w:val'),'1'); trPr.append(cs)
        bidi_=OxmlElement('w:bidi'); trPr.append(bidi_)
        for idx,(cell,(txt,bold_)) in enumerate(zip(row.cells,[(field,True),(value,header)])):
            cell.text=""
            p=cell.paragraphs[0]; pPr=p._p.get_or_add_pPr()
            pPr.append(OxmlElement("w:bidi"))
            jc=OxmlElement("w:jc"); jc.set(qn("w:val"),"right"); pPr.append(jc)
            vr=p.add_run(str(txt) if txt else "—")
            vr.font.size=Pt(10); vr.font.name="Times New Roman"; vr.font.bold=bold_
            if header: vr.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
            tc=cell._tc; tcP=tc.get_or_add_tcPr()
            shd=OxmlElement('w:shd'); shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto')
            if header:    shd.set(qn('w:fill'),'2D2926' if idx==0 else '8B7355')
            elif idx==0:  shd.set(qn('w:fill'),'F7F3EE')
            else:         shd.set(qn('w:fill'),'FFFFFF')
            tcP.append(shd)
            mg=OxmlElement('w:tcMar')
            for side in ['top','bottom','left','right']:
                m=OxmlElement(f'w:{side}'); m.set(qn('w:w'),'60'); m.set(qn('w:type'),'dxa'); mg.append(m)
            tcP.append(mg)

    # Header
    p_hdr=doc.add_paragraph(); p_hdr.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p_hdr.paragraph_format.space_after=Pt(4)
    if os.path.exists(LOGO_FILE):
        try: p_hdr.add_run().add_picture(LOGO_FILE, width=Inches(2.5))
        except: pass
    rt=p_hdr.add_run("\nقائمة أكنباك للسلوك الطفلي — CBCL/6-18\nالتقرير السريري")
    rt.font.name="Times New Roman"; rt.font.size=Pt(16); rt.font.bold=True
    rt.font.color.rgb=DARK_RGB

    p_sub=doc.add_paragraph(); p_sub.alignment=WD_ALIGN_PARAGRAPH.CENTER
    rts=p_sub.add_run("ASEBA · Achenbach & Rescorla (2001)")
    rts.font.size=Pt(9); rts.font.color.rgb=WARM_RGB; rts.font.name="Times New Roman"

    p_sep=doc.add_paragraph(); p_sep.paragraph_format.space_before=Pt(4); p_sep.paragraph_format.space_after=Pt(8)
    pPr=p_sep._p.get_or_add_pPr(); pBdr2=OxmlElement('w:pBdr'); b2=OxmlElement('w:bottom')
    b2.set(qn('w:val'),'single'); b2.set(qn('w:sz'),'6'); b2.set(qn('w:space'),'2'); b2.set(qn('w:color'),'8B7355')
    pBdr2.append(b2); pPr.append(pBdr2)

    # Demographics
    demo_tbl=make_table_2col()
    add_row(demo_tbl,"الحقل","البيانات",header=True)
    add_row(demo_tbl,"الطفل",child_name or "—")
    add_row(demo_tbl,"السن",str(age) or "—")
    add_row(demo_tbl,"النوع",gender or "—")
    add_row(demo_tbl,"المُقيِّم",rater or "—")
    add_row(demo_tbl,"التاريخ",date.today().strftime('%Y/%m/%d'))
    add_row(demo_tbl,"المقياس","CBCL/6-18 — 113 بنداً، مقياس 0-1-2")
    add_row(demo_tbl,"فترة التقييم","الستة أشهر الماضية")
    doc.add_paragraph().paragraph_format.space_after=Pt(6)

    # Score tables
    def add_score_section(title, keys):
        add_section_title(title)
        st=doc.add_table(rows=0,cols=4); st.style='Table Grid'
        try:
            tPr=st._tbl.tblPr; bv=OxmlElement('w:bidiVisual'); tPr.append(bv)
            tW=OxmlElement('w:tblW'); tW.set(qn('w:w'),'9026'); tW.set(qn('w:type'),'dxa'); tPr.append(tW)
            tg=OxmlElement('w:tblGrid')
            for w in [3800,1100,1100,3026]:
                gc=OxmlElement('w:gridCol'); gc.set(qn('w:w'),str(w)); tg.append(gc)
            st._tbl.insert(0,tg)
        except: pass
        hrow=st.add_row()
        trPr=hrow._tr.get_or_add_trPr(); bd=OxmlElement('w:bidi'); trPr.append(bd)
        for ci,htxt in enumerate(["المقياس","الخام","التائي","التصنيف"]):
            cell=hrow.cells[ci]; cell.text=""
            p=cell.paragraphs[0]; pPr=p._p.get_or_add_pPr()
            pPr.append(OxmlElement("w:bidi"))
            jc=OxmlElement("w:jc"); jc.set(qn("w:val"),"right"); pPr.append(jc)
            vr=p.add_run(htxt); vr.font.bold=True; vr.font.size=Pt(9.5)
            vr.font.name="Times New Roman"; vr.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
            tc=cell._tc; tcP=tc.get_or_add_tcPr()
            shd=OxmlElement('w:shd'); shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto')
            shd.set(qn('w:fill'),'2D2926'); tcP.append(shd)
        for i,k in enumerate(keys):
            s=scores[k]; t=s['t']
            fill_main="F7F3EE" if i%2==0 else "FFFFFF"
            t_fills={0:"D4EDDA",60:"FFF3CD",65:"FFE0B2",70:"FFCDD2"}
            t_fill="FFCDD2" if t>=70 else "FFE0B2" if t>=65 else "FFF3CD" if t>=60 else "D4EDDA"
            row=st.add_row()
            trPr2=row._tr.get_or_add_trPr(); bd2=OxmlElement('w:bidi'); trPr2.append(bd2)
            for ci,(txt,fill_) in enumerate([(SUBSCALES[k]['name_ar'],fill_main),
                                              (str(s['raw']),fill_main),
                                              (str(t),t_fill),
                                              (get_level_ar(t),fill_main)]):
                cell=row.cells[ci]; cell.text=""
                p=cell.paragraphs[0]; pPr=p._p.get_or_add_pPr()
                pPr.append(OxmlElement("w:bidi"))
                jc=OxmlElement("w:jc"); jc.set(qn("w:val"),"right"); pPr.append(jc)
                vr=p.add_run(txt); vr.font.size=Pt(9.5); vr.font.name="Times New Roman"
                vr.font.bold=(ci==2)
                tc=cell._tc; tcP=tc.get_or_add_tcPr()
                shd=OxmlElement('w:shd'); shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto')
                shd.set(qn('w:fill'),fill_); tcP.append(shd)
                mg=OxmlElement('w:tcMar')
                for side in ['top','bottom','left','right']:
                    m=OxmlElement(f'w:{side}'); m.set(qn('w:w'),'60'); m.set(qn('w:type'),'dxa'); mg.append(m)
                tcP.append(mg)
        doc.add_paragraph().paragraph_format.space_after=Pt(4)

    add_score_section("مقاييس المتلازمات السلوكية", ["ANX","WIT","SOM","SOC","THO","ATT","RUL","AGG"])
    add_score_section("المقاييس الشاملة", ["INT","EXT","TOT"])
    add_score_section("مقاييس موجهة DSM", ["DEP","ANX_D","SOM_D","ADHD_D","ODD","CON"])

    # Charts
    for chart_key, title in [
        ("syndrome",  "ملف الدرجات التائية — مقاييس المتلازمات"),
        ("broadband", "المقاييس الشاملة"),
        ("dsm",       "مقاييس DSM الموجهة"),
        ("pie",       "توزيع الاستجابات"),
    ]:
        if charts.get(chart_key):
            add_section_title(title)
            pc=doc.add_paragraph(); pc.alignment=WD_ALIGN_PARAGRAPH.CENTER
            pc.paragraph_format.space_after=Pt(6)
            w_=Inches(5.5) if chart_key in ("syndrome","dsm") else Inches(4.0)
            try: pc.add_run().add_picture(io.BytesIO(charts[chart_key]), width=w_)
            except: pass

    # Narrative
    add_section_title("التقرير السريري التفصيلي")
    sec_ar_pat=re.compile(r'^[١٢٣٤٥٦٧٨٩\d]+[\.،:]\s+[\u0600-\u06FF]')
    AR_HDRS={"قائمة أكنباك للسلوك الطفلي","ملخص سريري","نظرة عامة","التوصيات","الملخص",
             "ملف المشكلات","تحليل","تفسير","نقاط القوة","الانطباعات","البنود الحرجة",
             "إجمالي المشكلات"}
    in_table=False; current_table=None

    for line in report_text.split('\n'):
        ls=line.strip()
        if not ls:
            if in_table: in_table=False; current_table=None
            doc.add_paragraph().paragraph_format.space_after=Pt(2); continue
        upper=ls.upper()
        is_sec = sec_ar_pat.match(ls) or any(h in ls for h in AR_HDRS) or ls=="ملخص سريري"
        if is_sec:
            in_table=False; current_table=None
            add_section_title(ls); continue
        if ls.startswith('━') or ls.startswith('═'):
            in_table=False; current_table=None; continue
        if '|' in ls:
            parts=[p.strip() for p in ls.split('|') if p.strip()]
            if not parts: continue
            skip=[("field","value"),("subscale","raw"),("المقياس","الخام"),("الحقل","البيانات")]
            if len(parts)>=2 and (parts[0].strip('* ').lower(),parts[1].strip('* ').lower()) in skip: continue
            if not in_table or current_table is None:
                in_table=True; current_table=make_table_2col()
                add_row(current_table,"الحقل","التفاصيل",header=True)
            add_row(current_table,parts[0].strip('* '),' | '.join(parts[1:])); continue
        in_table=False; current_table=None
        add_para(ls, size=10.5, space_before=0, space_after=3)

    # Item responses table
    doc.add_page_break()
    add_section_title("جدول استجابات البنود الكاملة")
    add_para("مفتاح التقييم: 0 = غير صحيح  ·  1 = صحيح أحياناً  ·  2 = صحيح غالباً",
             size=8.5, color=WARM_RGB, space_after=6)

    rating_labels_ar={0:"0 — غير صحيح",1:"1 — أحياناً",2:"2 — غالباً"}
    rating_fill={0:"D4EDDA",1:"FFF3CD",2:"FFCDD2"}
    item_tbl=doc.add_table(rows=0,cols=3); item_tbl.style='Table Grid'
    try:
        tPr2=item_tbl._tbl.tblPr; bv2=OxmlElement('w:bidiVisual'); tPr2.append(bv2)
        tW2=OxmlElement('w:tblW'); tW2.set(qn('w:w'),'9026'); tW2.set(qn('w:type'),'dxa'); tPr2.append(tW2)
    except: pass
    hrow=item_tbl.add_row()
    for ci,htxt in enumerate(["#","البند","التقييم"]):
        cell=hrow.cells[ci]; cell.text=""
        p_=cell.paragraphs[0]; pPr_=p_._p.get_or_add_pPr()
        pPr_.append(OxmlElement("w:bidi"))
        jc_=OxmlElement("w:jc"); jc_.set(qn("w:val"),"right"); pPr_.append(jc_)
        r_=p_.add_run(htxt); r_.font.bold=True; r_.font.size=Pt(9); r_.font.name="Times New Roman"
        r_.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
        tc_=cell._tc; tcP_=tc_.get_or_add_tcPr()
        shd_=OxmlElement('w:shd'); shd_.set(qn('w:val'),'clear')
        shd_.set(qn('w:color'),'auto'); shd_.set(qn('w:fill'),'2D2926'); tcP_.append(shd_)
    for i,item_text in enumerate(ITEMS_AR):
        item_num=i+1; val=responses_dict.get(item_num,0)
        irow=item_tbl.add_row()
        fill_bg="F7F3EE" if i%2==0 else "FFFFFF"
        for ci,cell_txt in enumerate([str(item_num),item_text,rating_labels_ar[val]]):
            cell=irow.cells[ci]; cell.text=""
            p_=cell.paragraphs[0]; pPr_=p_._p.get_or_add_pPr()
            pPr_.append(OxmlElement("w:bidi"))
            jc_=OxmlElement("w:jc"); jc_.set(qn("w:val"),"right"); pPr_.append(jc_)
            r_=p_.add_run(cell_txt); r_.font.size=Pt(8.5); r_.font.name="Times New Roman"
            r_.font.bold=(ci==2); r_.font.color.rgb=RGBColor(0,0,0)
            tc_=cell._tc; tcP_=tc_.get_or_add_tcPr()
            shd_=OxmlElement('w:shd'); shd_.set(qn('w:val'),'clear'); shd_.set(qn('w:color'),'auto')
            shd_.set(qn('w:fill'), rating_fill[val] if ci==2 else fill_bg); tcP_.append(shd_)
            mg_=OxmlElement('w:tcMar')
            for side in ['top','bottom','left','right']:
                m_=OxmlElement(f'w:{side}'); m_.set(qn('w:w'),'50'); m_.set(qn('w:type'),'dxa'); mg_.append(m_)
            tcP_.append(mg_)

    add_para("هذا التقرير سري للغاية. الدرجات مبنية على تقييم الوالدين وتُفسَّر بالتزامن مع الحكم السريري. "
             "الدرجات التائية ≥65 تعتبر في النطاق الحدي؛ ≥70 في النطاق الإكلينيكي. "
             "© 2001 T.M. Achenbach, ASEBA, University of Vermont.",
             size=8.5, color=WARM_RGB, space_before=8)

    buf=io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf

# ══════════════════════════════════════════════════════════════
#  EMAIL
# ══════════════════════════════════════════════════════════════
def send_email(child_name, buf_pdf, buf_doc, fn_pdf, fn_doc, scores, recipient):
    date_str=date.today().strftime('%B %d, %Y')
    elevated=[(k,scores[k]["t"]) for k in SCALE_ORDER if scores[k]["t"]>=65]
    elev_html="".join(
        f"<tr><td style='padding:4px 0;color:#6B5B45;'>{SUBSCALES[k]['name_en']}</td>"
        f"<td><strong style='color:#D9534F;'>T={t}</strong></td></tr>"
        for k,t in elevated
    ) or "<tr><td colspan='2' style='color:#4CAF50;'>No scales elevated ≥ 65</td></tr>"

    msg=MIMEMultipart('mixed')
    msg['From']=GMAIL_USER; msg['To']=recipient
    msg['Subject']=f"[CBCL/6-18] {child_name} — {date_str}"
    body=f"""<html><body style="font-family:Georgia,serif;color:#1C1917;background:#F7F3EE;padding:20px;">
  <div style="max-width:560px;margin:0 auto;background:white;border:1px solid #DDD5C8;border-radius:4px;padding:28px;">
    <h2 style="font-weight:300;font-size:20px;color:#1C1917;margin-bottom:4px;">CBCL/6-18 Report</h2>
    <p style="color:#8B7355;font-size:11px;margin-top:0;text-transform:uppercase;letter-spacing:.08em;">
      Child Behavior Checklist — Achenbach & Rescorla (2001)</p>
    <hr style="border:none;border-top:1px solid #DDD5C8;margin:16px 0;">
    <table style="width:100%;font-size:13px;border-collapse:collapse;">
      <tr><td style="padding:5px 0;color:#8B7355;width:40%;">Child</td><td><strong>{child_name}</strong></td></tr>
      <tr><td style="padding:5px 0;color:#8B7355;">Date</td><td>{date_str}</td></tr>
    </table>
    <hr style="border:none;border-top:1px solid #DDD5C8;margin:16px 0;">
    <p style="font-size:12px;color:#8B7355;font-weight:bold;margin-bottom:6px;">Elevated Scales (T≥65)</p>
    <table style="width:100%;font-size:12px;border-collapse:collapse;">{elev_html}</table>
    <hr style="border:none;border-top:1px solid #DDD5C8;margin:16px 0;">
    <p style="font-size:12px;">📄 <strong>English Report (PDF)</strong> + 📝 <strong>Arabic Report (Word)</strong></p>
    <p style="font-size:10px;color:#8B7355;font-style:italic;">Confidential — for the treating clinician only.</p>
  </div></body></html>"""
    msg.attach(MIMEText(body,'html'))
    for buf_,fname_,mime_ in [
        (buf_pdf, fn_pdf, 'application/pdf'),
        (buf_doc, fn_doc, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'),
    ]:
        buf_.seek(0)
        part=MIMEBase('application', mime_.split('/')[-1])
        part=MIMEBase(*mime_.split('/'))
        part.set_payload(buf_.read()); encoders.encode_base64(part)
        part.add_header('Content-Disposition','attachment',filename=fname_)
        msg.attach(part)
    with smtplib.SMTP_SSL('smtp.gmail.com',465) as srv:
        srv.login(GMAIL_USER,GMAIL_PASS)
        srv.sendmail(GMAIL_USER,recipient,msg.as_string())

# ══════════════════════════════════════════════════════════════
#  PAGE CONFIG & CSS  — identical to Conners
# ══════════════════════════════════════════════════════════════
st.set_page_config(page_title="CBCL/6-18 Assessment", page_icon="📋", layout="wide")

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=EB+Garamond:ital,wght@0,400;0,500;0,600;1,400&family=Cairo:wght@400;600&display=swap');
html, body, [class*="css"] {
    font-family: 'EB Garamond', 'Cairo', Georgia, serif;
    background-color: #F7F3EE;
}
.stApp { background-color: #F7F3EE; }

.q-card {
    background: white; border: 1px solid #DDD5C8; border-radius: 3px;
    padding: 14px 18px 8px; margin-bottom: 10px;
    box-shadow: 0 1px 3px rgba(28,25,23,0.06);
}
.q-num {
    font-size: 10px; font-weight: 600; color: #8B7355;
    letter-spacing: .1em; text-transform: uppercase; margin-bottom: 4px;
}
.q-text { font-size: 15px; color: #1C1917; line-height: 1.5; margin-bottom: 10px; }

div[data-testid="stRadio"] > label { display: none; }
div[data-testid="stRadio"] > div {
    gap: 8px !important; flex-direction: row !important; flex-wrap: wrap !important;
}
div[data-testid="stRadio"] > div > label {
    background: #F7F3EE !important; border: 1.5px solid #DDD5C8 !important;
    border-radius: 2px !important; padding: 6px 16px !important;
    font-size: 12px !important; color: #8B7355 !important;
    font-family: 'EB Garamond', Georgia, serif !important;
    cursor: pointer !important; transition: all 0.12s !important;
}
div[data-testid="stRadio"] > div > label:has(input:checked) {
    background: #2D2926 !important; color: white !important;
    border-color: #2D2926 !important;
}

div[data-testid="stTextInput"] input,
div[data-testid="stSelectbox"] select {
    background: white !important; border: 1px solid #DDD5C8 !important;
    border-radius: 2px !important;
    font-family: 'EB Garamond', 'Cairo', Georgia, serif !important;
}

.stButton > button {
    background: #2D2926 !important; color: white !important;
    border: none !important; border-radius: 2px !important;
    font-size: .95rem !important; font-weight: 500 !important;
    font-family: 'EB Garamond', Georgia, serif !important;
    letter-spacing: .04em !important;
    box-shadow: none !important; transition: all 0.15s !important;
}
.stButton > button:hover { background: #8B7355 !important; }

.progress-wrap {
    background: #DDD5C8; height: 2px; border-radius: 2px; margin: 6px 0 16px 0;
}
.progress-fill { background: #8B7355; height: 2px; border-radius: 2px; transition: width 0.3s; }

.lang-btn-active { border-bottom: 2px solid #8B7355 !important; color: #1C1917 !important; }
</style>
""", unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════
#  HEADER
# ══════════════════════════════════════════════════════════════
col_logo, col_hdr = st.columns([1, 5])
with col_logo:
    if os.path.exists(LOGO_FILE): st.image(LOGO_FILE, width=110)
with col_hdr:
    st.markdown("""
    <div style="padding: 8px 0 0 0;">
        <div style="font-size:1.55rem;font-weight:500;color:#1C1917;
                    font-family:'EB Garamond',Georgia,serif;letter-spacing:.01em;line-height:1.2;">
            Child Behavior Checklist — CBCL/6-18
        </div>
        <div style="font-size:.8rem;color:#8B7355;margin-top:3px;letter-spacing:.08em;text-transform:uppercase;">
            Achenbach System of Empirically Based Assessment · Ages 6–18
        </div>
    </div>""", unsafe_allow_html=True)

st.markdown('<hr style="border:none;border-top:1px solid #DDD5C8;margin:12px 0 16px;">', unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════
#  SESSION STATE
# ══════════════════════════════════════════════════════════════
if "lang"        not in st.session_state: st.session_state.lang = "en"
if "responses"   not in st.session_state: st.session_state.responses = {}
if "report_done" not in st.session_state: st.session_state.report_done = False

# ══════════════════════════════════════════════════════════════
#  RESULTS SCREEN
# ══════════════════════════════════════════════════════════════
if st.session_state.report_done:
    scores     = st.session_state["scores"]
    rt_en      = st.session_state["report_en"]
    rt_ar      = st.session_state["report_ar"]
    child_name = st.session_state["child_name"]
    age        = st.session_state["child_age"]
    gender     = st.session_state["child_gender"]
    rater      = st.session_state["rater"]
    responses_v= st.session_state["responses_v"]
    lang       = st.session_state.lang

    int_t=scores["INT"]["t"]; ext_t=scores["EXT"]["t"]; tot_t=scores["TOT"]["t"]
    int_lbl = get_level_en(int_t) if lang=="en" else get_level_ar(int_t)
    ext_lbl = get_level_en(ext_t) if lang=="en" else get_level_ar(ext_t)
    tot_lbl = get_level_en(tot_t) if lang=="en" else get_level_ar(tot_t)

    st.markdown(f"""
    <div style="background:#1C1917;color:white;padding:16px 24px;border-radius:3px;margin-bottom:18px;">
        <span style="font-size:1.15rem;font-family:'EB Garamond',Georgia,serif;">
            ✓ &nbsp; <strong>{child_name}</strong> &nbsp;·&nbsp;
            Internalizing T={int_t} ({int_lbl}) &nbsp;·&nbsp;
            Externalizing T={ext_t} ({ext_lbl}) &nbsp;·&nbsp;
            Total T={tot_t} ({tot_lbl})
        </span>
    </div>""", unsafe_allow_html=True)

    # Metrics
    mc = st.columns(5)
    for i,(k,lbl) in enumerate([("INT","Internalizing"),("EXT","Externalizing"),("TOT","Total"),
                                  ("ATT","Attention"),("AGG","Aggressive")]):
        with mc[i]:
            t=scores[k]["t"]
            st.metric(lbl, f"T={t}", get_level_en(t) if lang=="en" else get_level_ar(t))

    # Charts
    bar_b    = make_syndrome_chart(scores)
    broad_b  = make_broadband_chart(scores)
    dsm_b    = make_dsm_chart(scores)
    pie_b    = make_pie_chart(responses_v)
    charts   = {"syndrome":bar_b,"broadband":broad_b,"dsm":dsm_b,"pie":pie_b}

    st.subheader("📊 Score Profiles" if lang=="en" else "📊 ملفات الدرجات")
    st.image(bar_b, use_container_width=True)
    c1,c2 = st.columns(2)
    with c1: st.image(broad_b, use_container_width=True)
    with c2: st.image(dsm_b,   use_container_width=True)
    with st.columns([1,2,1])[1]: st.image(pie_b, use_container_width=True)

    # Report tabs
    tab_en, tab_ar = st.tabs(["🇬🇧 English Report","🇸🇦 التقرير العربي"])
    with tab_en: st.text_area("",value=rt_en,height=500,label_visibility="collapsed")
    with tab_ar: st.text_area("",value=rt_ar,height=500,label_visibility="collapsed")

    # Downloads & email
    st.markdown('<hr style="border:none;border-top:1px solid #DDD5C8;margin:20px 0 12px;">', unsafe_allow_html=True)
    fn_pdf = f"{child_name.replace(' ','_')}_CBCL_EN.pdf"
    fn_doc = f"{child_name.replace(' ','_')}_CBCL_AR.docx"

    dl1,dl2,dl3,dl4,dl5 = st.columns(5)
    with dl1:
        buf_pdf=build_pdf_report_en(rt_en,scores,charts,child_name,age,gender,rater,responses_v)
        st.download_button("📄 English PDF",data=buf_pdf,file_name=fn_pdf,
                           mime="application/pdf",use_container_width=True)
    with dl2:
        buf_doc=build_word_report_ar(rt_ar,scores,charts,child_name,age,gender,rater,responses_v)
        st.download_button("📄 Arabic Word",data=buf_doc,file_name=fn_doc,
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                           use_container_width=True)
    with dl3:
        email_inp=st.text_input("📧 Email address",placeholder="clinic@example.com",
                                label_visibility="collapsed",key="email_send_inp")
    with dl4:
        if st.button("Send Email / إرسال",use_container_width=True):
            recipient=email_inp.strip() if email_inp.strip() else RECIPIENT_EMAIL
            try:
                buf_pdf2=build_pdf_report_en(rt_en,scores,charts,child_name,age,gender,rater,responses_v)
                buf_doc2=build_word_report_ar(rt_ar,scores,charts,child_name,age,gender,rater,responses_v)
                send_email(child_name,buf_pdf2,buf_doc2,fn_pdf,fn_doc,scores,recipient)
                st.success(f"✅ Sent to {recipient}")
            except Exception as e:
                st.error(f"Email error: {e}")
    with dl5:
        if st.button("↺ New Assessment",use_container_width=True):
            for k in list(st.session_state.keys()):
                if k!="lang": del st.session_state[k]
            st.session_state.responses={}; st.session_state.report_done=False
            st.rerun()
    st.stop()

# ══════════════════════════════════════════════════════════════
#  LANGUAGE TOGGLE
# ══════════════════════════════════════════════════════════════
lang = st.session_state.lang
c1, c2, c3 = st.columns([2, 2, 8])
with c1:
    if st.button("🇬🇧 English", use_container_width=True):
        st.session_state.lang="en"; st.session_state.responses={}
        st.session_state.report_done=False; st.rerun()
with c2:
    if st.button("🇸🇦 العربية", use_container_width=True):
        st.session_state.lang="ar"; st.session_state.responses={}
        st.session_state.report_done=False; st.rerun()

st.markdown(f"""
<div style="font-size:.72rem;color:#8B7355;letter-spacing:.07em;text-transform:uppercase;margin-bottom:18px;">
    {'🇬🇧 English mode active' if lang=='en' else '🇸🇦 النسخة العربية نشطة'}
</div>""", unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════
#  INTAKE FORM
# ══════════════════════════════════════════════════════════════
if lang == "en":
    st.markdown("""<div style="font-size:.72rem;font-weight:600;color:#8B7355;
        letter-spacing:.08em;text-transform:uppercase;margin-bottom:.8rem;
        padding-bottom:.4rem;border-bottom:1px solid #DDD5C8;">
        Child Information</div>""", unsafe_allow_html=True)
    c1,c2,c3 = st.columns(3)
    with c1:
        child_name=st.text_input("Child's Full Name",placeholder="First and Last Name",key="cn_inp")
        child_age =st.text_input("Age",placeholder="e.g. 10",key="ca_inp")
    with c2:
        child_gender=st.radio("Gender",["Male","Female"],key="cg_inp",horizontal=True)
        child_grade =st.text_input("School Grade",placeholder="e.g. Grade 4",key="cgr_inp")
    with c3:
        rater       =st.text_input("Rater Name (Parent/Caregiver)",placeholder="Name",key="rn_inp")
        relationship=st.selectbox("Relationship to Child",
                                  ["—","Biological Parent","Step Parent","Grandparent",
                                   "Foster Parent","Other"],key="rel_inp")
        relationship="" if relationship=="—" else relationship

    st.markdown("""<div style="background:white;border:1px solid #DDD5C8;border-radius:2px;
        padding:1rem 1.4rem;margin:1.2rem 0;font-size:.88rem;color:#1C1917;line-height:1.9;">
        <span style="color:#8B7355;font-weight:500;letter-spacing:.04em;font-size:.75rem;">INSTRUCTIONS</span><br>
        Below is a list of items that describe children and youths. For each item that describes
        your child <strong>now or within the past 6 months</strong>, please choose:<br>
        <strong>0</strong> = Not True (as far as you know) &nbsp;·&nbsp;
        <strong>1</strong> = Somewhat or Sometimes True &nbsp;·&nbsp;
        <strong>2</strong> = Very True or Often True
    </div>""", unsafe_allow_html=True)
    SCALE_OPTS=["0 — Not True","1 — Somewhat True","2 — Very True"]
    ITEMS=ITEMS_EN; item_label="Item"

else:
    st.markdown("""<div style="font-size:.72rem;font-weight:600;color:#8B7355;
        letter-spacing:.08em;text-transform:uppercase;margin-bottom:.8rem;
        padding-bottom:.4rem;border-bottom:1px solid #DDD5C8;direction:rtl;text-align:right;">
        بيانات الطفل</div>""", unsafe_allow_html=True)
    c1,c2,c3=st.columns(3)
    with c1:
        child_name=st.text_input("اسم الطفل كاملاً",placeholder="الاسم الأول والأخير",key="cn_inp")
        child_age =st.text_input("السن",placeholder="مثال: 10",key="ca_inp")
    with c2:
        child_gender=st.radio("النوع",["ذكر","أنثى"],key="cg_inp",horizontal=True)
        child_grade =st.text_input("الصف الدراسي",placeholder="مثال: الصف الرابع",key="cgr_inp")
    with c3:
        rater       =st.text_input("اسم المُقيِّم (ولي الأمر)",placeholder="الاسم",key="rn_inp")
        relationship=st.selectbox("صلة القرابة بالطفل",
                                  ["—","الأم","الأب","الجدة","الجد","وصي","أخرى"],key="rel_inp")
        relationship="" if relationship=="—" else relationship

    st.markdown("""<div style="background:white;border:1px solid #DDD5C8;border-radius:2px;
        padding:1rem 1.4rem;margin:1.2rem 0;font-size:.88rem;color:#1C1917;line-height:1.9;
        direction:rtl;text-align:right;">
        <span style="color:#8B7355;font-weight:500;letter-spacing:.04em;font-size:.75rem;">التعليمات</span><br>
        فيما يلي قائمة ببنود تصف الأطفال والشباب. لكل بند ينطبق على طفلك
        <strong>الآن أو خلال الستة أشهر الماضية</strong>، يرجى الاختيار:<br>
        <strong>0</strong> = غير صحيح &nbsp;·&nbsp;
        <strong>1</strong> = صحيح أحياناً &nbsp;·&nbsp;
        <strong>2</strong> = صحيح غالباً أو دائماً
    </div>""", unsafe_allow_html=True)
    SCALE_OPTS=["0 — غير صحيح","1 — أحياناً","2 — غالباً"]
    ITEMS=ITEMS_AR; item_label="بند"

# ══════════════════════════════════════════════════════════════
#  103 ITEMS
# ══════════════════════════════════════════════════════════════
responses=st.session_state.responses
all_answered=True
direction='rtl' if lang=='ar' else 'ltr'
align='right' if lang=='ar' else 'left'

for idx, item_text in enumerate(ITEMS):
    item_num=idx+1
    st.markdown(f"""<div class="q-card" style="direction:{direction};">
        <div class="q-num" style="text-align:{align};">{item_label} {item_num} / 103</div>
        <div class="q-text" style="text-align:{align};">{item_text}</div>
    </div>""", unsafe_allow_html=True)

    saved=responses.get(item_num)
    choice=st.radio(
        f"item_{item_num}", SCALE_OPTS,
        index=saved, key=f"resp_{item_num}",
        horizontal=True, label_visibility="collapsed"
    )
    if choice is None:
        all_answered=False
    else:
        val=int(choice[0])
        responses[item_num]=val
        st.session_state.responses[item_num]=val

# ══════════════════════════════════════════════════════════════
#  PROGRESS & SUBMIT
# ══════════════════════════════════════════════════════════════
answered_count=len([v for v in responses.values() if v is not None])
pct=int((answered_count/103)*100)
prog_text=f"{answered_count} of 103 answered" if lang=="en" else f"{answered_count} من 103 بنداً"

st.markdown(f"""
<div style="text-align:center;font-size:.78rem;color:#8B7355;
            letter-spacing:.06em;margin-top:1.5rem;">{prog_text}</div>
<div class="progress-wrap">
    <div class="progress-fill" style="width:{pct}%"></div>
</div>""", unsafe_allow_html=True)

if not all_answered and answered_count > 0:
    warn=("⚠ Please answer all 103 items before submitting." if lang=="en"
          else "⚠ يرجى الإجابة على جميع البنود الـ 103 قبل الإرسال.")
    st.markdown(f"""<div style="background:#FFF8F0;border-left:3px solid #E07B39;
        padding:1rem 1.2rem;border-radius:0 2px 2px 0;
        font-size:.88rem;color:#7A3D1A;margin:1rem 0;">{warn}</div>""",
        unsafe_allow_html=True)

st.markdown("<br>", unsafe_allow_html=True)
btn_label="✦ Generate Report" if lang=="en" else "✦ توليد التقرير"
col_btn,_=st.columns([2,3])
with col_btn:
    submit=st.button(btn_label, use_container_width=True, disabled=(answered_count<103))

if submit and answered_count==103:
    child_name_v  = child_name   or ("Child"  if lang=="en" else "الطفل")
    child_age_v   = child_age    or "—"
    rater_v       = rater        or ("Parent" if lang=="en" else "ولي الأمر")
    gender_v      = child_gender
    gender_v_en   = gender_v if lang=="en" else ("Male" if gender_v=="ذكر" else "Female")
    responses_v   = dict(st.session_state.responses)

    with st.spinner("⏳ Scoring and generating reports..." if lang=="en"
                    else "⏳ جاري الحساب وإنشاء التقارير..."):
        scores    = compute_scores(responses_v)
        report_en = generate_report_en(child_name_v, child_age_v, gender_v_en, rater_v, scores)
        report_ar = generate_report_ar(child_name_v, child_age_v, gender_v, rater_v, scores)

        st.session_state["scores"]      = scores
        st.session_state["report_en"]   = report_en
        st.session_state["report_ar"]   = report_ar
        st.session_state["child_name"]  = child_name_v
        st.session_state["child_age"]   = child_age_v
        st.session_state["child_gender"]= gender_v_en
        st.session_state["rater"]       = rater_v
        st.session_state["responses_v"] = responses_v
        st.session_state.report_done    = True
        st.rerun()
